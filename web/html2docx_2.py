import sys
import time
import qdarkgraystyle
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *

TIME_LIMIT = 100

class counter(QThread):
    """
    Runs a counter thread.
    """
    countChanged = pyqtSignal(int)

    def run(self):
        count = 0
        while count < TIME_LIMIT:
            count += 25
            time.sleep(0.5)
            self.countChanged.emit(count)

class MyWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setupUI()

    def setupUI(self):
        self.setGeometry(800, 350, 300, 200)
        self.setWindowTitle("HTML2WORD_BY_HIJACK")
        self.setWindowIcon(QIcon('data//icon.jpg'))

        self.pushButton1 = QPushButton("Open File")
        self.pushButton1.clicked.connect(self.pushButtonClicked1)
        self.pushButton2 = QPushButton("Save File")
        self.pushButton2.clicked.connect(self.pushButtonClicked2)
        self.label1 = QLabel()
        self.label2 = QLabel()
        self.pushButton3 = QPushButton("실행")
        self.pushButton3.clicked.connect(self.pushButtonClicked3)
        self.pushButton4 = QPushButton("종료")
        self.pushButton4.clicked.connect(QCoreApplication.instance().quit)
        self.pushButton5 = QPushButton("Open Folder")
        self.pushButton5.clicked.connect(self.pushButtonClicked4)
        self.label3 = QLabel()
        self.pushButton6 = QPushButton("전체 실행")
        self.pushButton6.clicked.connect(self.pushButtonClicked5)
        self.view = QListView()
        self.model = QStandardItemModel()
        self.label4 = QLabel('HTML2DOCX ver 0.2 by HIJACK ACADEMIA')
        myFont = QFont()
        myFont.setBold(True)
        self.label4.setFont(myFont)
        self.labeldoc = QLabel(doc)
        self.label4.setWordWrap(True)
        self.pbar = QProgressBar(self)
        self.pbar.setMaximum(100)
        pixmap = QPixmap('data//logo.png')
        self.label5 = QLabel()
        self.label5.setPixmap(QPixmap(pixmap))
        self.label5.setOpenExternalLinks(True)
        layout1 = QGridLayout()
        layout1.addWidget(self.pushButton1, 0, 0)
        layout1.addWidget(self.pushButton2, 1, 0)
        layout1.addWidget(self.label1, 0, 1)
        layout1.addWidget(self.label2, 1, 1)
        layout1.addWidget(self.pushButton3, 0, 2)
        layout1.addWidget(self.pushButton4, 1, 2)
        layout1.addWidget(self.pushButton5, 2, 0)
        layout1.addWidget(self.label3, 2, 1)
        layout1.addWidget(self.pushButton6, 2, 2)
        layout2 = QVBoxLayout()
        layout2.addWidget(self.view)
        layout3 = QVBoxLayout()
        layout3.addWidget(self.label4)
        layout3.addWidget(self.labeldoc)
        layout4 = QHBoxLayout()
        layout4.addWidget(self.pbar)
        layout4.addWidget(self.label5)
        layout = QVBoxLayout()
        layout.addLayout(layout1)
        layout.addLayout(layout2)
        layout.addLayout(layout3)
        layout.addLayout(layout4)
        self.setLayout(layout)


    def pushButtonClicked1(self):
        fname1 = QFileDialog.getOpenFileName(self)
        self.label1.setText(fname1[0])
        global fname0
        fname0 = fname1[0]
    def pushButtonClicked2(self):
        fname2 = QFileDialog.getSaveFileName(self)
        self.label2.setText(fname2[0])
        global sf
        sf = fname2[0]
    def pushButtonClicked3(self):
        # %% beuatiful soup4 & Python-Docx
        from docx import Document
        from bs4 import BeautifulSoup

        # %%
        document = Document('data//form.docx')

        # %%
        with open(str(fname0), encoding='UTF8') as fp:
            soup = BeautifulSoup(fp, 'html.parser')
        ex_id_test1 = soup.find_all(id='paragraph')
        # %%
        data1 = []
        for name in ex_id_test1:
            data1.append(name.get_text())

        # %%
        document.add_paragraph('- 원본지문 : ', style='second')
        document.add_paragraph('- 핵심소재 :', style='second')
        document.add_paragraph('- 핵심내용 :', style='third')

        # %%
        i = 0
        while i < len(data1):
            if (i % 2 == 0):
                document.add_paragraph('□ ' + data1[i], style='paragraph1')
                i = i + 1;
            else:
                document.add_paragraph('▶ ' + data1[i], style='paragraph2')
                i = i + 1;
        if 'docx' in str(sf):
            document.save(str(sf))
        else:
            document.save(str(sf)+'.docx')

            self.calc = counter()
            self.calc.countChanged.connect(self.onCountChanged)
            self.calc.start()

            QMessageBox.about(self, "message", "변환이 완료되었습니다.")

    def pushButtonClicked4(self):
        import os
        options = QFileDialog.Options()
        options |= QFileDialog.ShowDirsOnly
        global folderdirfirst
        folderdirfirst = QFileDialog.getExistingDirectory(self)
        self.label3.setText(folderdirfirst[0])
        global folderdir
        folderdir = folderdirfirst[0]
        for file in os.listdir(str(folderdirfirst)):
            if file.endswith(".html"):
                self.model.appendRow(QStandardItem(file))
        self.view.setModel(self.model)
    def pushButtonClicked5(self):
        # %% beuatiful soup4 & Python-Docx

        from docx import Document
        from bs4 import BeautifulSoup
        import os
        global files
        for file in os.listdir(str(folderdirfirst)):
            if file.endswith(".html"):
                files.append(file)
                document = Document('data//form.docx')
                with open(str(folderdirfirst)+'//'+str(file), encoding='UTF8') as fp:
                    soup = BeautifulSoup(fp, 'html.parser')
                ex_id_test1 = soup.find_all(id='paragraph')
                # %%
                data1 = []
                for name in ex_id_test1:
                   data1.append(name.get_text())

                # %%
                document.add_paragraph('- 원본지문 : ', style='second')
                document.add_paragraph('- 핵심소재 :', style='second')
                document.add_paragraph('- 핵심내용 :', style='third')

                # %%
                i = 0
                while i < len(data1):
                    if (i % 2 == 0):
                        document.add_paragraph('□ ' + data1[i], style='paragraph1')
                        i = i + 1;
                    else:
                        document.add_paragraph('▶ ' + data1[i], style='paragraph2')
                        i = i + 1;
                document.save(str(folderdirfirst)+'//'+str(file[:-5])+'.docx')

                self.calc = counter()
                self.calc.countChanged.connect(self.onCountChanged)
                self.calc.start()

        QMessageBox.about(self, "message", "변환이 완료되었습니다.")



    def onCountChanged(self, value):
        self.pbar.setValue(value)

        # %%


if __name__ == "__main__":
    fname0 = []
    sf = []
    folderdir = []
    files = []
    folderdirfirst = []
    doc = '하이잭 빅데이터 자동 워드파일 변환기입니다. \n 1. Open File 을 눌러 저장해 둔 HTML 파일을 선택해 주세요 \n 2. Save File 을 눌러 저장할 위치를 선택 후 파일이름을 설정해 주세요.\n 3. 변환 버튼을 누르시면 변환이 시작됩니다.'
    app = QApplication(sys.argv)
    window = MyWindow()
    app.setStyleSheet(qdarkgraystyle.load_stylesheet())
    window.show()
    app.exec_()